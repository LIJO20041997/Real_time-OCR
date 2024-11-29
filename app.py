import streamlit as st
import os
import tempfile
from pdf2image import convert_from_path
import cv2
import numpy as np
import easyocr
from docx import Document

def detect_tables_and_extract_text_to_docx_one_row_per_table(pdf_path, output_docx_path, output_image_dir):
    reader = easyocr.Reader(['en', 'hi'])
    pages = convert_from_path(pdf_path, dpi=300)
    os.makedirs(output_image_dir, exist_ok=True)
    doc = Document()

    for page_num, page_image in enumerate(pages, start=1):
        doc.add_heading(f"Page {page_num}", level=1)
        page_image_path = os.path.join(output_image_dir, f"page_{page_num}.png")
        page_image.save(page_image_path, "PNG")
        image = cv2.imread(page_image_path)
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        _, binary = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY_INV)
        horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (25, 1))
        vertical_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 25))
        horizontal_lines = cv2.morphologyEx(binary, cv2.MORPH_OPEN, horizontal_kernel, iterations=2)
        vertical_lines = cv2.morphologyEx(binary, cv2.MORPH_OPEN, vertical_kernel, iterations=2)
        table_structure = cv2.add(horizontal_lines, vertical_lines)
        contours, _ = cv2.findContours(table_structure, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

        for contour in contours:
            x, y, w, h = cv2.boundingRect(contour)
            if w > 50 and h > 50:
                table_img = image[y:y + h, x:x + w]
                result = reader.readtext(table_img)
                table_text = []
                for _, text, _ in result:
                    table_text.append(text)
                doc.add_paragraph(" | ".join(table_text))

        output_image_path = os.path.join(output_image_dir, f"page_{page_num}_detected.png")
        cv2.imwrite(output_image_path, image)

    doc.save(output_docx_path)

# Streamlit App
def main():
    st.title("Table Extraction from PDF to DOCX")
    st.write("Upload a PDF file, extract table data, and download it as a Word document.")

    # Initialize session state variables
    if "docx_ready" not in st.session_state:
        st.session_state.docx_ready = False
        st.session_state.output_docx_path = None

    uploaded_file = st.file_uploader("Upload a PDF file", type=["pdf"])
    if uploaded_file and not st.session_state.docx_ready:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
            temp_pdf.write(uploaded_file.read())
            temp_pdf_path = temp_pdf.name

        output_dir = tempfile.mkdtemp()
        st.session_state.output_docx_path = os.path.join(output_dir, "extracted_tables.docx")

        with st.spinner("Processing..."):
            detect_tables_and_extract_text_to_docx_one_row_per_table(
                temp_pdf_path, st.session_state.output_docx_path, output_dir
            )

        st.session_state.docx_ready = True
        st.success("Table extraction complete! You can now download the DOCX file.")

    if st.session_state.docx_ready and st.session_state.output_docx_path:
        with open(st.session_state.output_docx_path, "rb") as docx_file:
            st.download_button(
                label="Download Extracted Tables DOCX",
                data=docx_file,
                file_name="extracted_tables.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == "__main__":
    main()
