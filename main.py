from pdf2image import convert_from_path
import cv2
import numpy as np
import easyocr
import os
from docx import Document

def detect_tables_and_extract_text_to_docx_one_row_per_table(pdf_path, output_docx_path, output_image_dir):
    # Initialize EasyOCR reader
    reader = easyocr.Reader(['en', 'hi'])

    # Convert PDF to images
    pages = convert_from_path(pdf_path, dpi=300)
    
    # Ensure the output image directory exists
    os.makedirs(output_image_dir, exist_ok=True)

    # Create a Word document
    doc = Document()

    # Process each page
    for page_num, page_image in enumerate(pages, start=1):
        # Add a heading for the page
        doc.add_heading(f"Page {page_num}", level=1)

        # Save page image as a temporary file
        page_image_path = os.path.join(output_image_dir, f"page_{page_num}.png")
        page_image.save(page_image_path, "PNG")

        # Load the image
        image = cv2.imread(page_image_path)
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)

        # Apply binary thresholding
        _, binary = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY_INV)

        # Define kernels for morphological operations
        horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (25, 1))
        vertical_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 25))

        # Detect horizontal lines
        horizontal_lines = cv2.morphologyEx(binary, cv2.MORPH_OPEN, horizontal_kernel, iterations=2)

        # Detect vertical lines
        vertical_lines = cv2.morphologyEx(binary, cv2.MORPH_OPEN, vertical_kernel, iterations=2)

        # Combine horizontal and vertical lines
        table_structure = cv2.add(horizontal_lines, vertical_lines)

        # Find contours in the table structure
        contours, _ = cv2.findContours(table_structure, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

        # Draw bounding boxes around detected tables and extract text
        for contour in contours:
            x, y, w, h = cv2.boundingRect(contour)
            if w > 50 and h > 50:  # Filter out small boxes
                # Extract text using EasyOCR from the detected bounding box
                table_img = image[y:y + h, x:x + w]
                result = reader.readtext(table_img)

                # Combine all the table's rows into a single row
                table_text = []
                for _, text, _ in result:
                    table_text.append(text)

                # Add the combined row to the Word document
                doc.add_paragraph(" | ".join(table_text))  # Separate items with a pipe symbol for clarity

        # Save the modified image with bounding boxes
        output_image_path = os.path.join(output_image_dir, f"page_{page_num}_detected.png")
        cv2.imwrite(output_image_path, image)

    # Save the Word document
    doc.save(output_docx_path)

# Example usage
detect_tables_and_extract_text_to_docx_one_row_per_table(
    "sample.pdf",
    "table_data.docx",
    "output_images"
)
